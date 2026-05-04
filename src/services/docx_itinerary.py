import os
from datetime import datetime, timezone
from pathlib import Path

from docx import Document


def get_doc_output_dir() -> Path:
    raw = os.getenv("DOC_OUTPUT_DIR", "var/generated_docs")
    p = Path(raw)
    if not p.is_absolute():
        p = Path.cwd() / p
    p = p.resolve()
    p.mkdir(parents=True, exist_ok=True)
    return p


def _new_filename(session_id: int) -> str:
    ts = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%S")
    return f"itinerary_{session_id}_{ts}.docx"


def build_itinerary_docx(*, session_id: int, title: str, body_text: str) -> str:
    """
    在 DOC_OUTPUT_DIR 下写入 docx。返回 **仅 basename**，写入 DocSession.output_path。
    """
    out_dir = get_doc_output_dir()
    filename = _new_filename(session_id)
    full_path = out_dir / filename

    doc = Document()
    doc.add_heading(title, 0)
    for line in body_text.splitlines():
        doc.add_paragraph(line)
    # 纯空行时 splitlines 会丢连续空行，可接受于第一版
    doc.save(str(full_path))
    return filename


def resolve_output_file(stored: str) -> Path:
    """
    将库里存的 basename 安全解析为绝对路径；防路径穿越。
    """
    if (
        not stored
        or stored != Path(stored).name
        or ".." in stored
        or "/" in stored
        or "\\" in stored
    ):
        raise ValueError("非法的 output_path")
    out_dir = get_doc_output_dir()
    path = (out_dir / stored).resolve()
    if not path.is_file():
        raise FileNotFoundError
    if not path.is_relative_to(out_dir):
        raise ValueError("禁止访问输出目录外文件")
    return path
